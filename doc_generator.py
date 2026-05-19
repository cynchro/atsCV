from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle, KeepTogether
)


# ─── Palette ───────────────────────────────────────────────────────────────────
DARK       = "#0D1B2A"   # casi negro azulado — nombre y títulos
ACCENT     = "#1B6CA8"   # azul medio — línea decorativa y label de skills
BODY       = "#2C3E50"   # gris oscuro — cuerpo de texto
GRAY       = "#7F8C8D"   # gris medio — fechas, ubicación, contacto
RULE_LIGHT = "#BDC3C7"   # gris claro — líneas de sección

DARK_RGB   = RGBColor(0x0D, 0x1B, 0x2A)
ACCENT_RGB = RGBColor(0x1B, 0x6C, 0xA8)
GRAY_RGB   = RGBColor(0x7F, 0x8C, 0x8D)
BODY_RGB   = RGBColor(0x2C, 0x3E, 0x50)


# ══════════════════════════════════════════════════════════════════════════════
#  WORD (.docx)
# ══════════════════════════════════════════════════════════════════════════════

def _set_paragraph_space(p, before=0, after=0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _add_bottom_border(paragraph, color_hex="BDC3C7", size="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_top_border(paragraph, color_hex="1B6CA8", size="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color_hex)
    pBdr.append(top)
    pPr.append(pBdr)


def _labels(cv: dict) -> dict:
    defaults = {
        "summary": "Professional Summary",
        "experience": "Work Experience",
        "education": "Education",
        "skills": "Skills",
        "certifications": "Certifications",
        "skill_technical": "Technical",
        "skill_tools": "Tools",
        "skill_soft": "Soft Skills",
        "skill_languages": "Languages",
    }
    defaults.update(cv.get("section_labels", {}))
    return defaults


def generate_docx(cv: dict, output_path: str):
    doc = Document()

    # Page setup — tight margins for a full, professional look
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── NAME ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_space(p, before=0, after=2)
    run = p.add_run(cv.get("name", "").upper())
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_RGB

    # Thick accent line under the name
    accent_p = doc.add_paragraph()
    _set_paragraph_space(accent_p, before=0, after=4)
    _add_top_border(accent_p, color_hex="1B6CA8", size="18")

    # ── CONTACT ───────────────────────────────────────────────────────────────
    contact = cv.get("contact", {})
    parts = [v for v in [
        contact.get("email"), contact.get("phone"),
        contact.get("linkedin"), contact.get("location"),
        contact.get("website"),
    ] if v]
    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_space(p, before=0, after=10)
        run = p.add_run("  ·  ".join(parts))
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY_RGB

    # ── Helper: section header ─────────────────────────────────────────────────
    def section_header(title):
        p = doc.add_paragraph()
        _set_paragraph_space(p, before=10, after=1)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_RGB
        _add_bottom_border(p, color_hex="BDC3C7", size="4")

    # ── Helper: two-column row (left bold, right gray small) ──────────────────
    def two_col_row(left_bold: str, left_normal: str, right_text: str):
        """Job title (bold) — Company  |  dates (gray right side)"""
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        # Remove all borders
        for row in tbl.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    border = OxmlElement(f"w:{side}")
                    border.set(qn("w:val"), "none")
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        left_cell = tbl.rows[0].cells[0]
        right_cell = tbl.rows[0].cells[1]

        # Set column widths (approx 70/30)
        tbl.columns[0].width = Cm(11.5)
        tbl.columns[1].width = Cm(4.0)

        lp = left_cell.paragraphs[0]
        lp.clear()
        r1 = lp.add_run(left_bold)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = DARK_RGB
        if left_normal:
            r2 = lp.add_run(f"  —  {left_normal}")
            r2.font.size = Pt(10)
            r2.font.color.rgb = BODY_RGB

        rp = right_cell.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rp.add_run(right_text)
        rr.font.size = Pt(9)
        rr.font.color.rgb = GRAY_RGB

        _set_paragraph_space(lp, before=4, after=0)
        _set_paragraph_space(rp, before=4, after=0)

    def small_gray(text):
        p = doc.add_paragraph(text)
        _set_paragraph_space(p, before=0, after=2)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = GRAY_RGB

    lbl = _labels(cv)

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    if cv.get("summary"):
        section_header(lbl["summary"])
        p = doc.add_paragraph(cv["summary"])
        _set_paragraph_space(p, before=3, after=0)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = BODY_RGB

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    if cv.get("experience"):
        section_header(lbl["experience"])
        for job in cv["experience"]:
            date = f"{job.get('start_date', '')} – {job.get('end_date', '')}"
            two_col_row(job.get("title", ""), job.get("company", ""), date)

            if job.get("location"):
                small_gray(job["location"])

            for bullet in job.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(bullet).font.size = Pt(10)
                bp.runs[0].font.color.rgb = BODY_RGB
                bp.paragraph_format.left_indent = Cm(0.5)
                _set_paragraph_space(bp, before=1, after=1)

            doc.add_paragraph()  # breathing room between jobs

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    if cv.get("education"):
        section_header(lbl["education"])
        for edu in cv["education"]:
            degree = edu.get("degree", "")
            field  = edu.get("field", "")
            degree_str = f"{degree} in {field}" if field else degree
            two_col_row(degree_str, edu.get("institution", ""), edu.get("graduation_date", ""))

            if edu.get("location"):
                small_gray(edu["location"])
            extras = "  |  ".join(filter(None, [edu.get("gpa", ""), edu.get("honors", "")]))
            if extras:
                small_gray(extras)

    # ── SKILLS ────────────────────────────────────────────────────────────────
    skills = cv.get("skills", {})
    if any(skills.get(k) for k in ("technical", "tools", "soft", "languages")):
        section_header(lbl["skills"])
        mapping = [
            ("technical", lbl["skill_technical"]),
            ("tools",     lbl["skill_tools"]),
            ("soft",      lbl["skill_soft"]),
            ("languages", lbl["skill_languages"]),
        ]
        for key, label in mapping:
            items = skills.get(key, [])
            if items:
                p = doc.add_paragraph()
                _set_paragraph_space(p, before=2, after=1)
                r1 = p.add_run(f"{label}:  ")
                r1.bold = True
                r1.font.size = Pt(10)
                r1.font.color.rgb = ACCENT_RGB
                r2 = p.add_run(", ".join(items))
                r2.font.size = Pt(10)
                r2.font.color.rgb = BODY_RGB

    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    if cv.get("certifications"):
        section_header(lbl["certifications"])
        for cert in cv["certifications"]:
            p = doc.add_paragraph()
            _set_paragraph_space(p, before=2, after=1)
            r1 = p.add_run(cert.get("name", ""))
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = DARK_RGB
            suffix = "  —  ".join(filter(None, [cert.get("issuer", ""), cert.get("date", "")]))
            if suffix:
                r2 = p.add_run(f"  —  {suffix}")
                r2.font.size = Pt(9)
                r2.font.color.rgb = GRAY_RGB

    doc.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
#  PDF (ReportLab)
# ══════════════════════════════════════════════════════════════════════════════

PAGE_W = A4[0] - 4.4 * cm  # usable width after margins


def generate_pdf(cv: dict, output_path: str):
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2.2 * cm,
        leftMargin=2.2 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
    )

    C_DARK   = colors.HexColor(DARK)
    C_ACCENT = colors.HexColor(ACCENT)
    C_BODY   = colors.HexColor(BODY)
    C_GRAY   = colors.HexColor(GRAY)
    C_RULE   = colors.HexColor(RULE_LIGHT)

    # — Styles —
    name_s = ParagraphStyle(
        "name", fontSize=22, fontName="Helvetica-Bold",
        textColor=C_DARK, alignment=TA_LEFT, spaceAfter=6, spaceBefore=0,
    )
    contact_s = ParagraphStyle(
        "contact", fontSize=9, fontName="Helvetica",
        textColor=C_GRAY, alignment=TA_LEFT, spaceAfter=16, spaceBefore=6,
    )
    section_s = ParagraphStyle(
        "section", fontSize=10, fontName="Helvetica-Bold",
        textColor=C_DARK, spaceBefore=14, spaceAfter=2,
    )
    body_s = ParagraphStyle(
        "body", fontSize=10, fontName="Helvetica",
        textColor=C_BODY, spaceAfter=3, leading=14,
    )
    bullet_s = ParagraphStyle(
        "bullet", fontSize=10, fontName="Helvetica",
        textColor=C_BODY, leftIndent=14, spaceAfter=2, leading=14,
    )
    job_title_s = ParagraphStyle(
        "job_title", fontSize=10, fontName="Helvetica-Bold",
        textColor=C_DARK, spaceAfter=0, spaceBefore=8,
    )
    job_company_s = ParagraphStyle(
        "job_company", fontSize=10, fontName="Helvetica",
        textColor=C_BODY, spaceAfter=0, spaceBefore=8,
    )
    date_s = ParagraphStyle(
        "date", fontSize=9, fontName="Helvetica",
        textColor=C_GRAY, alignment=TA_RIGHT, spaceAfter=0, spaceBefore=8,
    )
    small_s = ParagraphStyle(
        "small", fontSize=9, fontName="Helvetica-Oblique",
        textColor=C_GRAY, spaceAfter=2,
    )
    skill_label_s = ParagraphStyle(
        "skill_label", fontSize=10, fontName="Helvetica-Bold",
        textColor=C_ACCENT, spaceAfter=0,
    )

    story = []

    # ── NAME ──────────────────────────────────────────────────────────────────
    story.append(Paragraph(cv.get("name", "").upper(), name_s))
    story.append(Spacer(1, 6))

    # Thick accent line under name
    story.append(HRFlowable(width="100%", thickness=3, color=C_ACCENT, spaceAfter=0, spaceBefore=0))
    story.append(Spacer(1, 8))

    # ── CONTACT ───────────────────────────────────────────────────────────────
    contact = cv.get("contact", {})
    parts = [v for v in [
        contact.get("email"), contact.get("phone"),
        contact.get("linkedin"), contact.get("location"),
        contact.get("website"),
    ] if v]
    if parts:
        story.append(Paragraph("  ·  ".join(parts), contact_s))
    story.append(Spacer(1, 10))

    # ── Helper: section header ─────────────────────────────────────────────────
    def add_section(title):
        story.append(Spacer(1, 4))
        story.append(Paragraph(title.upper(), section_s))
        story.append(HRFlowable(
            width="100%", thickness=0.75, color=C_RULE,
            spaceAfter=4, spaceBefore=1,
        ))

    # ── Helper: two-column row (title+company left, date right) ───────────────
    def job_row(title: str, company: str, date: str):
        left  = Paragraph(
            f"<b>{title}</b>  <font color='{BODY}'>{company}</font>",
            job_title_s,
        )
        right = Paragraph(date, date_s)
        t = Table([[left, right]], colWidths=[PAGE_W * 0.72, PAGE_W * 0.28])
        t.setStyle(TableStyle([
            ("VALIGN",       (0, 0), (-1, -1), "BOTTOM"),
            ("LEFTPADDING",  (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING",   (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 0),
        ]))
        story.append(t)

    def edu_row(degree_str: str, institution: str, date: str):
        left  = Paragraph(
            f"<b>{degree_str}</b>  <font color='{BODY}'>{institution}</font>",
            job_title_s,
        )
        right = Paragraph(date, date_s)
        t = Table([[left, right]], colWidths=[PAGE_W * 0.72, PAGE_W * 0.28])
        t.setStyle(TableStyle([
            ("VALIGN",       (0, 0), (-1, -1), "BOTTOM"),
            ("LEFTPADDING",  (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING",   (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 0),
        ]))
        story.append(t)

    lbl = _labels(cv)

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    if cv.get("summary"):
        add_section(lbl["summary"])
        story.append(Paragraph(cv["summary"], body_s))

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    if cv.get("experience"):
        add_section(lbl["experience"])
        for job in cv["experience"]:
            date = f"{job.get('start_date', '')} – {job.get('end_date', '')}"
            block = []
            block.append(Spacer(1, 2))
            # Build the job block as a KeepTogether so title doesn't orphan
            # (only keep header + first 2 bullets together)
            job_row(job.get("title", ""), job.get("company", ""), date)
            if job.get("location"):
                story.append(Paragraph(job["location"], small_s))
            for bullet in job.get("bullets", []):
                story.append(Paragraph(f"• {bullet}", bullet_s))
            story.append(Spacer(1, 6))

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    if cv.get("education"):
        add_section(lbl["education"])
        for edu in cv["education"]:
            degree = edu.get("degree", "")
            field  = edu.get("field", "")
            degree_str = f"{degree} in {field}" if field else degree
            edu_row(degree_str, edu.get("institution", ""), edu.get("graduation_date", ""))
            if edu.get("location"):
                story.append(Paragraph(edu["location"], small_s))
            extras = "  |  ".join(filter(None, [edu.get("gpa", ""), edu.get("honors", "")]))
            if extras:
                story.append(Paragraph(extras, small_s))
            story.append(Spacer(1, 5))

    # ── SKILLS ────────────────────────────────────────────────────────────────
    skills = cv.get("skills", {})
    if any(skills.get(k) for k in ("technical", "tools", "soft", "languages")):
        add_section(lbl["skills"])
        mapping = [
            ("technical", lbl["skill_technical"]),
            ("tools",     lbl["skill_tools"]),
            ("soft",      lbl["skill_soft"]),
            ("languages", lbl["skill_languages"]),
        ]
        for key, label in mapping:
            items = skills.get(key, [])
            if items:
                row_text = f"<font color='{ACCENT}'><b>{label}:</b></font>  {', '.join(items)}"
                story.append(Paragraph(row_text, body_s))
        story.append(Spacer(1, 4))

    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    if cv.get("certifications"):
        add_section(lbl["certifications"])
        for cert in cv["certifications"]:
            suffix = "  —  ".join(filter(None, [cert.get("issuer", ""), cert.get("date", "")]))
            line = f"<b>{cert.get('name', '')}</b>"
            if suffix:
                line += f"  <font size='9' color='{GRAY}'>{suffix}</font>"
            story.append(Paragraph(line, body_s))

    doc.build(story)
